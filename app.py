import streamlit as st
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

import tempfile
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

pdfmetrics.registerFont(TTFont('Arial', 'Arial.ttf'))

# ===== STYLE PDF =====
styles = getSampleStyleSheet()

styles.add(ParagraphStyle(
    name='UnitTitle',
    fontName='Arial',
    fontSize=18,
    alignment=1,
    spaceAfter=12))

styles.add(ParagraphStyle(
    name='VocabTitle',
    fontName='Arial',
    fontSize=14,
    spaceAfter=6))

styles.add(ParagraphStyle(
    name='DotLine',
    fontName='Arial',
    fontSize=13,
    leading=19.5))

# ===== HÃ m táº¡o dÃ²ng cháº¥m PDF =====
def dot_groups_pdf(word, per_line, space_count):
    clean_word = word.replace(" ", "")
    length = len(clean_word) * 3
    one_group = "." * length
    spaces = "&nbsp;" * space_count  # giá»¯ nguyÃªn khoáº£ng tráº¯ng
    return spaces.join([one_group] * per_line)

# ===== HÃ m táº¡o dÃ²ng cháº¥m WORD =====
def dot_groups_word(word, per_line, space_count):
    clean_word = word.replace(" ", "")
    length = len(clean_word) * 3
    one_group = "." * length
    spaces = "\u00A0" * space_count  # non-breaking space
    return spaces.join([one_group] * per_line)

# ===== Block PDF =====
def vocab_block_pdf(word, meaning, lines, per_line, space_count):
    elements = []
    title = f"<b>{word}: {meaning}</b>"
    elements.append(Paragraph(title, styles['VocabTitle']))

    for _ in range(lines):
        elements.append(Paragraph(dot_groups_pdf(word, per_line, space_count), styles['DotLine']))

    elements.append(Spacer(1, 12))
    return elements

# ===== Block WORD =====
def vocab_block_word(doc, word, meaning, lines, per_line, space_count):
    p = doc.add_paragraph()
    run = p.add_run(f"{word}: {meaning}")
    run.bold = True

    for _ in range(lines):
        p = doc.add_paragraph(dot_groups_word(word, per_line, space_count))
        p.paragraph_format.line_spacing = 1.5
        for r in p.runs:
            r.font.size = Pt(13)

# ===== Giao diá»‡n =====
st.title("ðŸ“˜ Táº¡o File Luyá»‡n Viáº¿t Tá»« Vá»±ng Cho BÃ©")

unit_name = st.text_input("TÃªn Unit")
num_words = st.number_input("Sá»‘ tá»« vá»±ng", min_value=1, step=1)

vocab_list = []

for i in range(int(num_words)):
    st.markdown(f"### Tá»« {i+1}")
    col1, col2 = st.columns(2)
    eng = col1.text_input(f"Tiáº¿ng Anh {i+1}")
    vie = col2.text_input(f"NghÄ©a {i+1}")

    col3, col4 = st.columns(2)
    lines = col3.number_input(f"Sá»‘ dÃ²ng viáº¿t tá»« {i+1}", min_value=1, step=1, value=3)
    per_line = col4.number_input(f"Má»—i dÃ²ng cÃ³ máº¥y láº§n viáº¿t tá»« {i+1}", min_value=1, step=1, value=3)

    space_count = st.number_input(f"Sá»‘ khoáº£ng tráº¯ng giá»¯a cÃ¡c cá»¥m cá»§a tá»« {i+1}", min_value=1, step=1, value=5)

    vocab_list.append((eng, vie, lines, per_line, space_count))

# ===== Táº O FILE =====
if st.button("ðŸ“„ Táº¡o file PDF & Word"):
    # ===== PDF =====
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_pdf:
        doc_pdf = SimpleDocTemplate(tmp_pdf.name, pagesize=A4,
                                    rightMargin=2*cm, leftMargin=2*cm,
                                    topMargin=2*cm, bottomMargin=2*cm)

        story = []
        story.append(Paragraph(f"Unit: {unit_name}", styles['UnitTitle']))
        story.append(Paragraph("<b>VOCABULARY</b>", styles['Heading2']))
        story.append(Spacer(1, 12))

        for word, meaning, lines, per_line, space_count in vocab_list:
            if word.strip():
                story.extend(vocab_block_pdf(word, meaning, lines, per_line, space_count))

        doc_pdf.build(story)

    # ===== WORD =====
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_word:
        doc_word = Document()

        # Unit Title
        title = doc_word.add_paragraph(f"Unit: {unit_name}")
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.runs[0]
        run.bold = True
        run.font.size = Pt(18)

        # Vocabulary heading
        vocab_head = doc_word.add_paragraph("VOCABULARY")
        vocab_head.runs[0].bold = True

        for word, meaning, lines, per_line, space_count in vocab_list:
            if word.strip():
                vocab_block_word(doc_word, word, meaning, lines, per_line, space_count)

        doc_word.save(tmp_word.name)

    # ===== Download buttons =====
    with open(tmp_pdf.name, "rb") as f:
        st.download_button("â¬‡ Táº£i PDF", f, file_name=f"{unit_name}_CopyWriting.pdf")

    with open(tmp_word.name, "rb") as f:
        st.download_button("â¬‡ Táº£i Word (.docx)", f, file_name=f"{unit_name}_CopyWriting.docx")